from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_THEME_COLOR
from docx.enum.style import WD_STYLE
from datetime import datetime
from datetime import timedelta
import pandas as pd
import matplotlib.pyplot as plt
from win32com.client import DispatchEx
from pathlib import Path
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_centered_para(text, font, size):
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_text(text)
    fnt = run.font
    fnt.color.theme_color = MSO_THEME_COLOR.ACCENT_1
    fnt.name = font
    fnt.size = Pt(size)
    return paragraph

def add_rightaligned_para(text, font, size):
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    run.add_text(text)
    fnt = run.font
    fnt.color.theme_color = MSO_THEME_COLOR.ACCENT_1
    fnt.name = font
    fnt.size = Pt(size)
    return paragraph
    
def add_leftaligned_para(text, font, size):
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()
    run.add_text(text)
    fnt = run.font
    fnt.name = font
    fnt.size = Pt(size)
    return paragraph

def add_new_heading(text):
    heading = document.add_heading('', 1).add_run()
    heading.add_text(text)
    heading.font.name = 'Arial'
    heading.font.size = Pt(16)
    heading.font.bold = True
    heading.font.color.theme_color = MSO_THEME_COLOR.ACCENT_1

def add_bullet(text):
    bullet = add_leftaligned_para(text, 'Arial', 12)
    bullet.runs[0].font.bold = True
    bullet.style = 'List Bullet'
    bullet.paragraph_format.space_after = Pt(0)

def add_new_table(rows, cols):
    table = document.add_table(rows=rows, cols=cols)
    table.style = 'Medium Shading 1 Accent 1'
    for i in range(cols):
        cell = table.cell(0,i)
        run = cell.add_paragraph().add_run()
        run.font.name = 'Arial'
        run.font.size = 14
    return table

def add_perday_table(df):
    t = add_new_table(2, 7)
    dates = pd.date_range(start=startdate.strftime("%Y-%m-%d"), end=datetime.now().strftime("%Y-%m-%d"))
    for i in range(7):
        t.cell(0,i).text = dates[i].strftime('%a')
        t.cell(1,i).text = str(df[df.Date == dates[i]].count())
        t.cell(0,i).paragraphs[0].runs[0].font.size = Pt(14)
        t.cell(1,i).paragraphs[0].runs[0].font.size = Pt(22)
        t.cell(1,i).paragraphs[0].runs[0].font.bold = True

        for j in range(2):
            t.cell(j,i).paragraphs[0].runs[0].font.name = "Arial"
            t.cell(j,i).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_event_table(df):
    t = add_new_table(df.shape[0]+1, df.shape[1])
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
        t.cell(0,j).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = t.cell(0,j).paragraphs[0].runs[0]
        run.font.size = Pt(14)
        run.font.name = 'Arial'

    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])
            t.cell(i+1,j).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = t.cell(i+1,j).paragraphs[0].runs[0]
            run.font.size = Pt(11)
            run.font.name = 'Arial'
            run.font.bold = False

def add_indented_para(text):
    indented = add_leftaligned_para(text, 'Arial', 11.5)
    indented.paragraph_format.left_indent = Inches(0.25)
    return indented

def pull_table(sheet, start_date, Unit):
    df = pd.read_excel(open(r'C:\Path\to\Excel Sheet.xlsx', 'rb'), engine='openpyxl', sheet_name=sheet)
    df = df.drop(labels=['Time', 'Unnamed: 6', 'Unnamed: 7'], axis=1, errors='ignore')
    mask = (df['Unit2 or Unit4'] == Unit)
    df = df.loc[mask]
    df['Date'] = pd.to_datetime(df['Date'], format="%Y-%m-%d")
    mask = (df['Date'] >= start_date - timedelta(days=1))
    df = df.loc[mask]
    df.insert(0, 'Date2', df['Date'].dt.date)
    df = df.drop(labels=['Date'], axis=1)
    df = df.rename(columns={"Date2":"Date"})
    df = df.drop(labels=['Unit2 or Unit4'], axis=1)
    return df

def add_space(num):
    for i in range(num):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)


def add_summary(sum):
    for i in sum:
        add_bullet(i[0])
        para = add_indented_para(i[1])
        add_metric(para, i[2])

def add_metric(para, num):
    para.add_run(num)
    para.runs[1].font.bold = True
    para.runs[1].font.name = 'Arial'
    para.runs[1].font.size = Pt(11.5)
    para.add_run('.')
    para.runs[2].font.name = 'Arial'
    para.runs[2].font.size = Pt(11.5)

def add_pie(values, labels, colors):
    plt.pie(pieValues, labels=labels, colors=colors,  startangle = 90)
    plt.title('Events Processed this Week')
    plt.savefig('pie.png', bbox_inches = 'tight')
    document.add_picture('pie.png', width=Inches(6.5))

Unit = ['Unit2', 'Unit4']
for h in Unit:
    document = Document()
    document.sections[0].top_margin = Inches(0.8)
    document.sections[0].bottom_margin = Inches(0.8)
    document.sections[0].right_margin = Inches(1)
    document.sections[0].left_margin = Inches(1)

    # Title Page
    document.add_picture('Logo.png', width=Inches(6.5))

    add_centered_para('Business Name', 'Arial', 28)
    add_centered_para(h, 'Arial', 28)
    add_space(4)
    add_centered_para('WEEKLY REPORT', 'Arial', 22)

    startdate = datetime.now() - timedelta(days=6)
    daterange = startdate.strftime("%B %d, %Y") + ' - ' + datetime.now().strftime("%B %d, %Y")
    add_centered_para(daterange, 'Arial', 16)

    add_space(11)
    add_rightaligned_para('REVIEWED BY:__________ DATE:_________', 'Arial', 10)
    footer = document.sections[0].footer
    footerrun = footer.paragraphs[0].add_run()
    footerrun.add_text('CONFIDENTIAL')
    footerrun.font.name = 'Arial'
    footerrun.font.size = Pt(12)
    #document.add_page_break()

    # Table of Contents
    document.add_section()
    footer2 = document.sections[1].footer
    footer2.is_linked_to_previous = False
    footer.paragraphs[0].text = "CONFIDENTIAL"
    
    #footerrun2.add_text('\tPage 2')
    footerrun2 = footer.paragraphs[0].runs[0]
    footerrun2.font.name = 'Arial'
    footerrun2.font.size = Pt(12)

    add_centered_para('TABLE OF CONTENTS', 'Arial', 16)
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "0-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p

    document.add_page_break()
    
    # Executive Summary
    add_new_heading('EXECUTIVE SUMMARY')
    add_leftaligned_para('This is a weekly summary of the events processed and policies created by Braintrace in Unit2 and Unit4. This summary is broken down into three sections with a daily total and list of policies / events.', 'Arial', 11.5)

    policies = pull_table('Number of New Policies', startdate, h)
    escalated = pull_table('Number of Escalated Events', startdate, h)
    solved = pull_table('Events Solved by UEP', startdate, h)

    summary = [['New Policies Created', 
        'These are policies created by Braintrace to clean up events. These were created after the analyst reviewed existing policies and determined a new one was required. The total number of new policies created this week was ',
        str(len(policies.index))],
        ['Escalated Events',
        'These are events that Braintrace determined needed to be escalated for additional review. The total number of escalated events this week was ',
        str(len(escalated.index))],
        ['Events Solved by Unenforced Policies (UEP)',
        'Some events can be resolved with modifying existing policies. Number of events solved by older policies this week was ',
        str(len(solved.index))]]

    add_summary(summary)
   
    pieValues = [len(policies.index), len(escalated.index), len(solved.index)]
    total = sum(pieValues)
    labels = ['New Policies ' + "{:.0%}".format(pieValues[0]/total), 'Escalated Events ' + "{:.0%}".format(pieValues[1]/total), 'UEP ' + "{:.0%}".format(pieValues[2]/total)]
    colors = ['#58585B',  '#3B8989', '#4472C4']
    add_pie(pieValues, labels, colors)

    document.add_page_break()
    
    # New Policies Created
    add_new_heading('NEW POLICIES CREATED')
    add_leftaligned_para('Below is a daily table of the number of new policies Braintrace has created this week.', 'Arial', 11.5)

    add_perday_table(policies)
    add_space(1)

    add_metric(add_leftaligned_para('The total number of new policies created this week is ', 'Arial', 11.5), str(len(policies.index)))
    add_leftaligned_para('The new policies created can be found in the table below.', 'Arial', 11.5)
    add_space(1)

    add_event_table(policies)
    document.add_page_break()

    # Events Escalated
    add_new_heading('EVENTS ESCALATED')
    add_leftaligned_para('Below is a daily table of the number of events Braintrace has escalated this week.', 'Arial', 11.5)

    add_perday_table(escalated)
    add_space(1)

    add_metric(add_leftaligned_para('The total number of events that were escalated this week is ', 'Arial', 11.5), str(len(escalated.index)))
    add_leftaligned_para('The events escalated can be found in the table below.', 'Arial', 11.5)
    add_space(1)

    add_event_table(escalated)
    document.add_page_break()

    # Solved by UEP
    add_new_heading('EVENTS SOLVED BY UNENFORCED POLICIES (UEP)')
    add_leftaligned_para('Below is a daily table of the number of events that were solved by older polices.', 'Arial', 11.5)

    add_perday_table(solved)
    add_space(1)

    add_metric(add_leftaligned_para('The total number of events that were solved by UEP this week is ', 'Arial', 11.5), str(len(solved.index)))
    add_leftaligned_para('The events escalated can be found in the table below.', 'Arial', 11.5)
    add_space(1)

    add_event_table(solved)
    
    filename = h + ".docx"
    document.save(filename)

    out_file = repr(Path.cwd() / filename)
    print(out_file)
    WordDoc = DispatchEx("Word.Application")
    #WordDoc.Visible = True
    WordDoc = WordDoc.Documents.Open(r"C:\Path\to\Business Name Insight Report\\" + filename)
    WordDoc.Sections(2).Footers(1).PageNumbers.Add(2,True)
    WordDoc.Sections(2).Footers(1).PageNumbers.NumberStyle = 15

    WordDoc.TablesOfContents(1).Update()
    WordDoc.Close()    
    